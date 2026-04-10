from pathlib import Path
from docx import Document

class DocxBuilder:
    def __init__(self, out_dir: str):
        self.out_dir = Path(out_dir)
        self.out_dir.mkdir(parents=True, exist_ok=True)

    def create_manual(self, title: str, sections):
        doc = Document()
        doc.add_heading(title, 0)
        for section in sections:
            doc.add_heading(section["title"], level=1)
            doc.add_paragraph(section["body"])
        path = self.out_dir / (title.replace(" ", "_") + ".docx")
        doc.save(path)
        return str(path)
